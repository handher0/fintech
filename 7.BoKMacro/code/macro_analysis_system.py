"""
거시경제 종합 분석 및 보고서 자동 생성 시스템

1단계(macro_data) KPI·챗봇 + 2단계(analyze_correlation) 상관관계 분석을 통합하고,
버튼 클릭 시 DOCX 종합 보고서를 생성한다.

실행 방법:
  pip install streamlit plotly matplotlib seaborn python-docx python-dotenv openai pandas requests numpy
  cd 7.BoKMacro
  streamlit run code/macro_analysis_system.py
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openai import OpenAI

from analyze_correlation import (
    ANALYSIS_START_D,
    ANALYSIS_START_M,
    SELECTABLE_COLUMNS,
    build_analysis_context,
    build_dual_axis,
    build_heatmap,
    build_master_dataframe,
    build_scatter,
    correlation_matrix,
    detect_chat_charts,
    strongest_pair,
    summary_stats,
)
from macro_data import (
    BASE_DIR,
    CHARTS_DIR,
    CPI_FETCH_START_M,
    INDICATOR_CONFIGS,
    MACRO_SYSTEM_PROMPT,
    _end_date_for_cycle,
    _strip_markdown,
    build_data_context,
    build_plotly_chart,
    call_llm,
    detect_chart_indicator,
    fetch_indicator,
    filter_display_period,
    kpi_cpi_yoy,
    kpi_exchange,
    kpi_rate,
    format_as_of,
    load_merged_env,
    needs_global_web_search,
    resolve_api_keys,
    run_web_search,
    show_key_error,
)

REPORTS_DIR = BASE_DIR / "output" / "reports"
FONT_NAME = "맑은 고딕"

UNIFIED_SYSTEM_PROMPT = MACRO_SYSTEM_PROMPT + """

추가 역할:
- 결합된 월별 데이터와 피어슨·스피어만 상관계수를 함께 참고하여 거시경제·자산 영향을 해석하세요.
- 상관관계를 인과관계로 단정하지 마세요.
"""

WEB_SEARCH_QUERY = (
    "2026년 현재 미국 연준(Fed) 금리 방향 및 글로벌 인플레이션 원인에 대해 "
    "핵심 수치를 서두에 명시하여 간략히 정리해 주세요."
)


# ---------------------------------------------------------------------------
# DOCX 한글 서식
# ---------------------------------------------------------------------------
def _format_korean_date(dt: datetime) -> str:
    return f"{dt.year}년{dt.month:02d}월{dt.day:02d}일"


def _set_run_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _disable_auto_spacing(paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "0")
        ppr.append(el)


def _add_korean_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    _disable_auto_spacing(p)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _add_cover_page(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    for line, sz, bold in (
        ("거시경제 종합 분석 보고서", 22, True),
        ("Macro Economic Analysis Report", 14, False),
        (_format_korean_date(datetime.now()), 12, False),
    ):
        p = cell.paragraphs[0] if line == "거시경제 종합 분석 보고서" else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _disable_auto_spacing(p)
        run = p.add_run(line)
        _set_run_font(run, size=sz, bold=bold)
    doc.add_page_break()


def _add_corr_table(doc: Document, corr: pd.DataFrame, title: str) -> None:
    _add_korean_paragraph(doc, title, bold=True, size=13)
    if corr.empty:
        _add_korean_paragraph(doc, "(상관계수 데이터 없음)")
        return
    labels = list(corr.columns)
    table = doc.add_table(rows=len(labels) + 1, cols=len(labels) + 1)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = ""
    for j, col in enumerate(labels, start=1):
        table.rows[0].cells[j].text = col
    for i, row_label in enumerate(labels, start=1):
        table.rows[i].cells[0].text = row_label
        for j, col_label in enumerate(labels, start=1):
            val = corr.iloc[i - 1, j - 1]
            table.rows[i].cells[j].text = "" if pd.isna(val) else f"{val:.2f}"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _disable_auto_spacing(p)
                for run in p.runs:
                    _set_run_font(run, size=10)
    doc.add_paragraph("")


def _add_stats_table(doc: Document, stats: pd.DataFrame) -> None:
    _add_korean_paragraph(doc, "지표별 통계 요약", bold=True, size=13)
    if stats.empty:
        _add_korean_paragraph(doc, "(통계 데이터 없음)")
        return
    cols = list(stats.columns)
    table = doc.add_table(rows=len(stats) + 1, cols=len(cols))
    table.style = "Table Grid"
    for j, c in enumerate(cols):
        table.rows[0].cells[j].text = str(c)
    for i, row in stats.iterrows():
        for j, c in enumerate(cols):
            table.rows[i + 1].cells[j].text = str(row[c])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _disable_auto_spacing(p)
                for run in p.runs:
                    _set_run_font(run, size=10)
    doc.add_paragraph("")


def _add_figure(doc: Document, path: str, caption: str) -> None:
    if not Path(path).exists():
        return
    _add_korean_paragraph(doc, caption, bold=True, size=11)
    doc.add_picture(path, width=Inches(5.5))
    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# 리포트용 정적 차트 (Matplotlib, kaleido 미사용)
# ---------------------------------------------------------------------------
def _save_report_charts(
    macro: dict[str, Any],
    master: pd.DataFrame,
    pearson: pd.DataFrame,
    ts: str,
) -> list[tuple[str, str]]:
    """시계열·히트맵·산점도 PNG를 저장하고 (경로, 한글 캡션) 목록을 반환한다."""
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    saved: list[tuple[str, str]] = []
    plt.rcParams["axes.unicode_minus"] = False

    try:
        ex = filter_display_period(macro["exchange"])
        if not ex.empty:
            path = CHARTS_DIR / f"report_exchange_{ts}.png"
            plt.figure(figsize=(10, 5))
            sns.lineplot(data=ex, x="DATE", y="VALUE", marker="o")
            plt.title("원/달러 환율 추이")
            ax = plt.gca()
            ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=8))
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
            plt.xticks(rotation=30, ha="right")
            plt.tight_layout()
            plt.savefig(path, dpi=150, bbox_inches="tight")
            plt.close()
            saved.append((str(path), "그림 1. 원/달러 환율 추이"))

        if not pearson.empty:
            path = CHARTS_DIR / f"report_heatmap_{ts}.png"
            plt.figure(figsize=(7, 5))
            sns.heatmap(pearson, annot=True, fmt=".2f", cmap="RdBu_r", center=0, vmin=-1, vmax=1)
            plt.title("지표 간 상관계수 히트맵")
            plt.tight_layout()
            plt.savefig(path, dpi=150, bbox_inches="tight")
            plt.close()
            saved.append((str(path), "그림 2. 지표 간 상관계수 히트맵"))

        plot_df = master[["base_rate", "exchange_avg"]].dropna()
        if len(plot_df) >= 2:
            path = CHARTS_DIR / f"report_scatter_{ts}.png"
            x = plot_df["base_rate"].values.astype(float)
            y = plot_df["exchange_avg"].values.astype(float)
            slope, intercept = np.polyfit(x, y, 1)
            r = float(np.corrcoef(x, y)[0, 1])
            plt.figure(figsize=(7, 5))
            plt.scatter(x, y, alpha=0.7)
            x_line = np.linspace(x.min(), x.max(), 50)
            plt.plot(x_line, slope * x_line + intercept, "r--", label=f"R={r:.2f}")
            plt.xlabel("기준금리")
            plt.ylabel("원/달러 환율(월평균)")
            plt.title("기준금리 vs 환율 산점도")
            plt.legend()
            plt.tight_layout()
            plt.savefig(path, dpi=150, bbox_inches="tight")
            plt.close()
            saved.append((str(path), "그림 3. 기준금리 vs 원/달러 환율 산점도"))
    except Exception:
        plt.close("all")

    return saved


# ---------------------------------------------------------------------------
# 리포트 생성
# ---------------------------------------------------------------------------
def generate_macro_report(
    client: OpenAI,
    model: str,
    macro: dict[str, Any],
    master: pd.DataFrame,
    pearson: pd.DataFrame,
    spearman: pd.DataFrame,
    use_web: bool,
) -> tuple[bytes, str]:
    """DOCX 종합 보고서를 생성하고 (bytes, 파일명)을 반환한다."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"macro_report_{ts}.docx"
    filepath = REPORTS_DIR / filename

    stats = summary_stats(master)
    pair = strongest_pair(pearson)
    chart_items = _save_report_charts(macro, master, pearson, ts)

    web_text = ""
    if use_web:
        try:
            web_text = run_web_search(client, WEB_SEARCH_QUERY)
        except Exception as exc:
            web_text = f"(웹 검색 실패: {exc})"

    ex_kpi = kpi_exchange(macro["exchange"])
    rt_kpi = kpi_rate(macro["rate"])
    cp_kpi = kpi_cpi_yoy(macro["cpi"])

    doc = Document()
    _add_cover_page(doc)

    _add_korean_paragraph(doc, "요약", bold=True, size=14)
    summary_lines = [
        f"분석 기간: {master['month'].iloc[0]} ~ {master['month'].iloc[-1]} (총 {len(master)}개월)",
        f"현재 원/달러 환율: {ex_kpi['value']:,.2f}원" if ex_kpi["value"] else "환율: 데이터 없음",
        f"한국은행 기준금리: {rt_kpi['value']:.2f}%" if rt_kpi["value"] else "기준금리: 데이터 없음",
        f"CPI 전년동월비: {cp_kpi['value']:.1f}%" if cp_kpi["value"] else "CPI YoY: 데이터 없음",
    ]
    if pair:
        summary_lines.append(f"피어슨 최강 상관 쌍: {pair[0]} ↔ {pair[1]} ({pair[2]:.3f})")
    for line in summary_lines:
        _add_korean_paragraph(doc, f"- {line}")

    _add_korean_paragraph(doc, "데이터 분석 통계", bold=True, size=14)
    _add_stats_table(doc, stats)
    _add_corr_table(doc, pearson.round(2), "피어슨 상관계수 행렬")
    _add_corr_table(doc, spearman.round(2), "스피어만 상관계수 행렬")

    _add_korean_paragraph(doc, "시각화", bold=True, size=14)
    for path, caption in chart_items:
        _add_figure(doc, path, caption)

    _add_korean_paragraph(doc, "글로벌 거시경제 원인 분석", bold=True, size=14)
    _add_korean_paragraph(doc, web_text or "(웹 검색 미사용 또는 결과 없음)")

    _add_korean_paragraph(doc, "종합 자산운용 결론", bold=True, size=14)
    conclusion_ctx = build_analysis_context(master, pearson, spearman)
    try:
        conclusion = call_llm(
            client, model,
            "위 통계를 바탕으로 투자자 관점의 종합 자산운용 결론을 5문장 이내로 작성하세요.",
            conclusion_ctx,
        )
    except Exception as exc:
        conclusion = (
            f"현재 환율·금리·물가 지표는 상관행렬과 시계열 추이를 함께 모니터링하며 "
            f"분산 투자와 유동성 관리를 병행할 것을 권고합니다. (LLM 오류: {exc})"
        )
    _add_korean_paragraph(doc, conclusion)

    doc.save(str(filepath))
    return filepath.read_bytes(), filename


# ---------------------------------------------------------------------------
# 데이터 로드
# ---------------------------------------------------------------------------
@st.cache_data(ttl=3600, show_spinner="거시경제 데이터 로드 중...")
def load_pooled_macro_data(api_key: str) -> dict[str, Any]:
    """2025년~ 수집, KPI·챗봇용 raw 시계열."""
    end_d = _end_date_for_cycle("D")
    end_m = _end_date_for_cycle("M")
    exchange_df, ex_cycle = fetch_indicator(
        api_key, INDICATOR_CONFIGS["exchange"],
        ANALYSIS_START_D, end_d, ANALYSIS_START_M, end_m,
    )
    rate_df, rate_cycle = fetch_indicator(
        api_key, INDICATOR_CONFIGS["rate"],
        ANALYSIS_START_D, end_d, ANALYSIS_START_M, end_m,
    )
    cpi_df, cpi_cycle = fetch_indicator(
        api_key, INDICATOR_CONFIGS["cpi"],
        ANALYSIS_START_D, end_d, CPI_FETCH_START_M, end_m,
    )
    return {
        "exchange": exchange_df,
        "exchange_cycle": ex_cycle,
        "rate": rate_df,
        "rate_cycle": rate_cycle,
        "cpi": cpi_df,
        "cpi_cycle": cpi_cycle,
    }


def build_unified_context(
    macro: dict[str, Any],
    master: pd.DataFrame,
    pearson: pd.DataFrame,
    spearman: pd.DataFrame,
) -> str:
    macro_ctx = build_data_context(macro["exchange"], macro["rate"], macro["cpi"])
    corr_ctx = build_analysis_context(master, pearson, spearman)
    return f"{macro_ctx}\n\n[상관관계 분석]\n{corr_ctx}"


# ---------------------------------------------------------------------------
# UI — KPI
# ---------------------------------------------------------------------------
def render_kpi_row(macro: dict[str, Any]) -> None:
    ex_kpi = kpi_exchange(macro["exchange"])
    rt_kpi = kpi_rate(macro["rate"])
    cp_kpi = kpi_cpi_yoy(macro["cpi"])
    c1, c2, c3 = st.columns(3)
    with c1:
        if ex_kpi["value"] is not None:
            st.metric("현재 원/달러 환율", f"{ex_kpi['value']:,.2f}원",
                      delta=f"{ex_kpi['delta']:+.2f}원" if ex_kpi["delta"] is not None else None)
            st.caption(format_as_of(ex_kpi["as_of"], "D"))
        else:
            st.metric("현재 원/달러 환율", "데이터 없음")
    with c2:
        if rt_kpi["value"] is not None:
            st.metric("한국은행 기준금리", f"{rt_kpi['value']:.2f}%",
                      delta=f"{rt_kpi['delta']:+.2f}%p" if rt_kpi["delta"] is not None else None)
            st.caption(format_as_of(rt_kpi["as_of"], macro["rate_cycle"]))
        else:
            st.metric("한국은행 기준금리", "데이터 없음")
    with c3:
        if cp_kpi["value"] is not None:
            st.metric("소비자물가 전년동월비", f"{cp_kpi['value']:.1f}%",
                      delta=f"{cp_kpi['delta']:+.2f}%p" if cp_kpi["delta"] is not None else None)
            st.caption(format_as_of(cp_kpi["as_of"], "M"))
        else:
            st.metric("소비자물가 전년동월비", "데이터 없음")


# ---------------------------------------------------------------------------
# UI — Tab 1: AI 매크로 에이전트
# ---------------------------------------------------------------------------
def render_agent_tab(
    macro: dict[str, Any],
    master: pd.DataFrame,
    pearson: pd.DataFrame,
    spearman: pd.DataFrame,
    client: OpenAI,
    model: str,
    use_web: bool,
    corr_method: str,
) -> None:
    if "messages" not in st.session_state:
        st.session_state.messages = []

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            for fig in msg.get("figs", []):
                st.plotly_chart(fig, use_container_width=True)

    prompt = st.chat_input("거시경제 지표 및 자산 영향에 대해 질문하세요.")
    if not prompt:
        return

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    context = build_unified_context(macro, master, pearson, spearman)
    web_context = ""
    if use_web and needs_global_web_search(prompt):
        with st.spinner("글로벌 경제 정보 검색 중..."):
            try:
                web_context = run_web_search(client, prompt)
            except Exception as exc:
                web_context = f"(웹 검색 오류: {exc})"

    figs: list[Any] = []
    with st.chat_message("assistant"):
        with st.spinner("분석 중..."):
            try:
                extra = f"\n\n[웹 검색 참고]\n{web_context}" if web_context else ""
                user_msg = f"[거시경제·상관관계 데이터]\n{context}{extra}\n\n[질문]\n{prompt}"
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": UNIFIED_SYSTEM_PROMPT},
                        {"role": "user", "content": user_msg},
                    ],
                    temperature=0.2,
                )
                answer = _strip_markdown((resp.choices[0].message.content or "").strip())
            except Exception as exc:
                answer = f"분석 중 오류가 발생했습니다: {exc}"
        st.markdown(answer)

        ind = detect_chart_indicator(prompt)
        if ind:
            cycle_key = f"{ind}_cycle" if ind != "cpi" else "cpi_cycle"
            fig = build_plotly_chart(ind, macro[ind], macro.get(cycle_key, "M"))
            if fig:
                figs.append(fig)
                st.plotly_chart(fig, use_container_width=True)

        hints = detect_chat_charts(prompt)
        active = pearson if corr_method == "pearson" else spearman
        method_label = "피어슨" if corr_method == "pearson" else "스피어만"
        if hints["heatmap"] and not active.empty:
            hm = build_heatmap(active, method_label)
            if hm:
                figs.append(hm)
                st.plotly_chart(hm, use_container_width=True)
        if hints["dual"]:
            dual = build_dual_axis(master, hints["col_a"], hints["col_b"])
            if dual:
                figs.append(dual)
                st.plotly_chart(dual, use_container_width=True)
        if hints["scatter"]:
            sc = build_scatter(master, hints["col_a"], hints["col_b"])
            if sc:
                figs.append(sc)
                st.plotly_chart(sc, use_container_width=True)

    st.session_state.messages.append({"role": "assistant", "content": answer, "figs": figs})


# ---------------------------------------------------------------------------
# UI — Tab 2: 상관관계
# ---------------------------------------------------------------------------
def render_correlation_tab(
    master: pd.DataFrame,
    corr_method: str,
) -> None:
    active = correlation_matrix(master, corr_method)
    method_label = "피어슨" if corr_method == "pearson" else "스피어만"

    st.subheader("분석 데이터 요약")
    st.metric("총 분석 개월 수", f"{len(master)}개월")
    st.dataframe(summary_stats(master), use_container_width=True, hide_index=True)

    st.subheader("1. 상관계수 히트맵")
    hm = build_heatmap(active, method_label)
    if hm:
        st.plotly_chart(hm, use_container_width=True)

    st.subheader("2. 이중 축 시계열 복합 대조")
    cols = list(SELECTABLE_COLUMNS.keys())
    c1, c2 = st.columns(2)
    with c1:
        sel_a = st.selectbox("좌측 Y축", cols, index=cols.index("base_rate"), key="sys_dual_a")
    with c2:
        sel_b = st.selectbox("우측 Y축", cols, index=cols.index("exchange_avg"), key="sys_dual_b")
    dual = build_dual_axis(master, sel_a, sel_b)
    if dual:
        st.plotly_chart(dual, use_container_width=True)

    st.subheader("3. 산점도 및 OLS 추세선")
    c3, c4 = st.columns(2)
    with c3:
        sel_x = st.selectbox("X축", cols, index=cols.index("base_rate"), key="sys_sc_x")
    with c4:
        sel_y = st.selectbox("Y축", cols, index=cols.index("exchange_avg"), key="sys_sc_y")
    sc = build_scatter(master, sel_x, sel_y)
    if sc:
        st.plotly_chart(sc, use_container_width=True)


# ---------------------------------------------------------------------------
# 세션
# ---------------------------------------------------------------------------
def init_session() -> None:
    defaults = {
        "messages": [],
        "docx_bytes": None,
        "docx_filename": None,
        "docx_version": 0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def reset_all() -> None:
    st.session_state.messages = []
    st.session_state.docx_bytes = None
    st.session_state.docx_filename = None
    st.session_state.docx_version = st.session_state.get("docx_version", 0) + 1


# ---------------------------------------------------------------------------
# 메인
# ---------------------------------------------------------------------------
def main() -> None:
    st.set_page_config(
        page_title="거시경제 종합 분석",
        page_icon="📊",
        layout="wide",
    )
    init_session()
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    merged_env, checked_paths = load_merged_env()
    ecos_key, openai_key, missing = resolve_api_keys(merged_env)

    st.title("거시경제 종합 분석 및 보고서 자동 생성 시스템")
    st.caption(
        "데이터: 2025년 1월~현재 풀링 · 시각화·챗봇 주 경계: 2026년 이후"
    )

    with st.sidebar:
        st.header("제어 패널")
        model = st.selectbox("LLM 모델", ["gpt-4o", "gpt-4-turbo", "gpt-4o-mini"], index=0)
        use_web = st.checkbox("OpenAI 실시간 글로벌 금융 트렌드 웹 검색 (web_search)", value=False)
        corr_method = st.radio(
            "상관계수 알고리즘",
            ["pearson", "spearman"],
            format_func=lambda x: "피어슨" if x == "pearson" else "스피어만",
        )
        if st.button("대화 초기화 (Reset Chat)", use_container_width=True):
            reset_all()
            st.rerun()
        st.divider()
        if st.button("종합 자산 보고서(DOCX) 생성", use_container_width=True, type="primary"):
            st.session_state.trigger_report = True

    if missing:
        show_key_error(missing, checked_paths)
        return

    try:
        client = OpenAI(api_key=openai_key)
    except Exception as exc:
        st.error(f"OpenAI 클라이언트 초기화 실패: {exc}")
        return

    try:
        macro = load_pooled_macro_data(ecos_key)
        master = build_master_dataframe(ecos_key)
    except Exception as exc:
        st.error(f"ECOS 데이터 수집 중 오류가 발생했습니다: {exc}")
        return

    if master.empty or len(master) < 3:
        st.error("결합된 분석 데이터가 부족합니다.")
        return

    pearson = correlation_matrix(master, "pearson")
    spearman = correlation_matrix(master, "spearman")

    render_kpi_row(macro)
    st.divider()

    tab_agent, tab_corr = st.tabs(["AI 매크로 에이전트", "통계적 자산 상관관계"])
    with tab_agent:
        render_agent_tab(macro, master, pearson, spearman, client, model, use_web, corr_method)
    with tab_corr:
        render_correlation_tab(master, corr_method)

    if st.session_state.get("trigger_report"):
        st.session_state.trigger_report = False
        with st.spinner("종합 보고서(DOCX) 생성 중..."):
            try:
                docx_bytes, filename = generate_macro_report(
                    client, model, macro, master, pearson, spearman, use_web
                )
                st.session_state.docx_bytes = docx_bytes
                st.session_state.docx_filename = filename
                st.session_state.docx_version = st.session_state.get("docx_version", 0) + 1
                st.success(f"보고서가 생성되었습니다: {filename}")
            except Exception as exc:
                st.error(f"보고서 생성 오류: {exc}")

    if st.session_state.get("docx_bytes"):
        st.divider()
        st.download_button(
            "종합 자산 보고서 다운로드 (DOCX)",
            data=st.session_state.docx_bytes,
            file_name=st.session_state.docx_filename or "macro_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_docx_{st.session_state.docx_version}",
        )


if __name__ == "__main__":
    main()
