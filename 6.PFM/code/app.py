"""
개인 맞춤형 자산 관리 가계부 Agent

기본 모드: 자연어 질문 -> LLM SQL 변환 -> SQLite 실행 -> Plotly 시각화 -> 잔소리 브리핑
리포트 모드: 버튼 클릭 시 -> 통계 + 웹 트렌드 + Matplotlib PNG -> python-docx 보고서

데이터: ./6.PFM/data/user_card_history_merged.csv (읽기 전용 마스터)
  메모리 내 SQLite(user_card_history)로 바인딩하여 SQL 실행

실행 방법:
  pip install streamlit plotly matplotlib seaborn python-docx python-dotenv openai pandas
  cd 6.PFM/code
  streamlit run app.py
"""

from __future__ import annotations

import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import plotly.express as px
import seaborn as sns
import streamlit as st
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI

# ---------------------------------------------------------------------------
# 경로
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parents[1]
CSV_PATH = BASE_DIR / "data" / "user_card_history_merged.csv"
CHARTS_DIR = BASE_DIR / "output" / "charts"
REPORTS_DIR = BASE_DIR / "output" / "reports"
REPORT_PATH = REPORTS_DIR / "PFM_Analysis_Report.docx"

TABLE_NAME = "user_card_history"
DB_COLUMNS = [
    "month", "approved_num", "approved_date_time",
    "card_name", "store_name", "category", "amount_krw",
]

SCHEMA_INFO = """
테이블명: user_card_history
컬럼:
  - month (TEXT) : 결제 월 (YYYY-MM, 2026-01 ~ 2026-05)
  - approved_num (TEXT) : 카드 승인번호
  - approved_date_time (TEXT) : 결제 일시 (YYYY-MM-DD HH:MM:SS)
  - card_name (TEXT) : 카드명 (신한 딥드림, 네이버페이 현대카드, 토스뱅크 체크카드)
  - store_name (TEXT) : 가맹점명 (배달의민족, 스타벅스, 카카오T(택시) 등)
  - category (TEXT) : 소비 분류 (식비/배달, 카페/간식, 교통/차량, 쇼핑/뷰티, 문화/생활)
  - amount_krw (REAL) : 결제 금액(원)
"""

SQL_SYSTEM_PROMPT = f"""당신은 SQLite SQL 전문가입니다. 아래 가계부 스키마를 참고해 사용자 질문에 맞는 SELECT 쿼리를 작성하세요.

{SCHEMA_INFO}

규칙:
- 오직 SELECT(또는 WITH ... SELECT)만 작성합니다.
- SQLite 문법을 사용합니다.
- 시간대 분석이 필요하면 strftime('%H', approved_date_time)를 사용하세요.
- 응답에는 SQL 쿼리만 출력하고 다른 설명은 하지 마세요.
"""

ANSWER_SYSTEM_PROMPT = """당신은 사용자의 소비를 진단하는 친근한 자산 관리 코치입니다.
SQL 실행 결과를 바탕으로 한국어로 답변하세요.

답변 형식 규칙 (반드시 준수):
- 마크다운 기호(#, ##, **, `, >, ---)를 절대 사용하지 마세요.
- 제목도 일반 문단 텍스트로 표기하세요.
- 수치 요약은 순수 문자열 불렛으로만 표기하세요. 예: "- 식비 지출: 99만 원"
- 줄바꿈과 순수 텍스트 위주로 출력하세요.
- 마지막에 소비 통제를 돕는 잔소리 한두 문장을 덧붙이세요.
"""


# ---------------------------------------------------------------------------
# 환경 변수 & 데이터
# ---------------------------------------------------------------------------
def load_api_key() -> str | None:
    """fintech/.env 기준으로 OPENAI_API_KEY를 로드한다."""
    repo_root = BASE_DIR.parent
    load_dotenv(repo_root / ".env", override=False)
    key = os.getenv("OPENAI_API_KEY", "").strip()
    return key or None


@st.cache_data(show_spinner=False)
def load_dataframe() -> pd.DataFrame:
    """마스터 CSV를 로드하고 정제해 반환한다."""
    df = pd.read_csv(CSV_PATH)
    if "month" in df.columns:
        df = df[df["month"].astype(str).str.strip().str.lower() != "total"]
    df["Amount_KRW"] = pd.to_numeric(df["Amount_KRW"], errors="coerce")
    df = df[df["Amount_KRW"].notna() & (df["Amount_KRW"] > 0)]
    return df.reset_index(drop=True)


@st.cache_resource(show_spinner=False)
def get_connection() -> sqlite3.Connection:
    """DataFrame을 메모리 SQLite로 바인딩한 연결을 반환한다."""
    df = load_dataframe().rename(
        columns={
            "month": "month",
            "Approved_Num": "approved_num",
            "Approved_DateTime": "approved_date_time",
            "Card_Name": "card_name",
            "Store_Name": "store_name",
            "Category": "category",
            "Amount_KRW": "amount_krw",
        }
    )
    conn = sqlite3.connect(":memory:", check_same_thread=False)
    df[DB_COLUMNS].to_sql(TABLE_NAME, conn, if_exists="replace", index=False)
    return conn


def is_safe_select(sql: str) -> bool:
    """SELECT/WITH 전용 쿼리인지 검증한다."""
    cleaned = re.sub(r"--.*$", "", sql, flags=re.MULTILINE)
    cleaned = re.sub(r"/\*.*?\*/", "", cleaned, flags=re.DOTALL).strip().upper()
    if not re.match(r"^\s*(WITH|SELECT)\b", cleaned):
        return False
    forbidden = r"\b(INSERT|UPDATE|DELETE|DROP|CREATE|ALTER|TRUNCATE|ATTACH|DETACH)\b"
    return re.search(forbidden, cleaned) is None


def extract_sql(text: str) -> str:
    """LLM 응답에서 SQL을 추출한다."""
    block = re.search(r"```(?:sql)?\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if block:
        return block.group(1).strip().rstrip(";")
    match = re.search(r"((?:WITH|SELECT)\b[\s\S]*)", text, re.IGNORECASE)
    if match:
        return match.group(1).strip().rstrip(";")
    return text.strip().rstrip(";")


def execute_sql(sql: str) -> pd.DataFrame:
    """검증된 SELECT를 실행한다."""
    if not is_safe_select(sql):
        raise ValueError("안전하지 않은 SQL입니다. 조회(SELECT)만 실행할 수 있습니다.")
    return pd.read_sql_query(sql, get_connection())


# ---------------------------------------------------------------------------
# LLM
# ---------------------------------------------------------------------------
def call_llm(client: OpenAI, model: str, system: str, user: str) -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=0.2,
    )
    return (resp.choices[0].message.content or "").strip()


def generate_sql(client: OpenAI, model: str, question: str) -> str:
    return extract_sql(call_llm(client, model, SQL_SYSTEM_PROMPT, question))


def summarize(client: OpenAI, model: str, question: str, sql: str, df: pd.DataFrame) -> str:
    preview = df.head(50).to_string(index=False) if not df.empty else "(결과 없음)"
    user_msg = (
        f"[사용자 질문]\n{question}\n\n"
        f"[실행 SQL]\n{sql}\n\n"
        f"[결과 행 수] {len(df)}건\n"
        f"[결과 미리보기]\n{preview}"
    )
    return call_llm(client, model, ANSWER_SYSTEM_PROMPT, user_msg)


def run_web_search(client: OpenAI, query: str) -> str:
    """OpenAI Responses API web_search로 금융 트렌드를 조회한다 (방어적)."""
    if not hasattr(client, "responses"):
        return ""
    try:
        resp = client.responses.create(
            model="gpt-5-nano",
            input=f"최신 개인 소비/가계부 절약 트렌드 관점에서 간략히 정리: {query}",
            tools=[{"type": "web_search"}],
        )
        if getattr(resp, "output_text", ""):
            return str(resp.output_text)
        parts = []
        for item in getattr(resp, "output", []) or []:
            if getattr(item, "type", None) == "message":
                for block in getattr(item, "content", []) or []:
                    if getattr(block, "type", None) == "output_text":
                        parts.append(getattr(block, "text", "") or "")
        return "\n".join(parts).strip()
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# 차트
# ---------------------------------------------------------------------------
def infer_chart_type(df: pd.DataFrame) -> str:
    if df.empty or len(df.columns) < 2:
        return "none"
    cols = [c.lower() for c in df.columns]
    if any("month" in c or "date" in c or "hour" in c for c in cols):
        return "line"
    if any(k in c for c in cols for k in ("category", "card", "store")):
        return "bar"
    if len(df.select_dtypes(include="number").columns) >= 1 and len(df) > 1:
        return "bar"
    return "none"


def build_plotly(df: pd.DataFrame, chart_type: str):
    if chart_type == "none" or df.empty:
        return None
    numeric = df.select_dtypes(include="number").columns.tolist()
    if not numeric:
        return None
    y = numeric[0]
    x = df.columns[0] if df.columns[0] != y else (df.columns[1] if len(df.columns) > 1 else df.columns[0])
    if chart_type == "line":
        fig = px.line(df, x=x, y=y, markers=True, title="추이 분석")
    else:
        fig = px.bar(df, x=x, y=y, title="비교 분석", color=x)
        fig.update_layout(showlegend=False)
    return fig


def save_png(df: pd.DataFrame, chart_type: str, chart_id: str) -> str | None:
    """리포트용 정적 PNG를 Matplotlib/Seaborn으로 저장한다 (kaleido 미사용)."""
    if chart_type == "none" or df.empty:
        return None
    numeric = df.select_dtypes(include="number").columns.tolist()
    if not numeric:
        return None
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    path = CHARTS_DIR / f"chart_{chart_id}.png"
    y = numeric[0]
    x = df.columns[0] if df.columns[0] != y else df.columns[1]
    try:
        plt.figure(figsize=(9, 5))
        sns.set_style("whitegrid")
        plt.rcParams["axes.unicode_minus"] = False
        if chart_type == "line":
            sns.lineplot(data=df, x=x, y=y, marker="o")
        else:
            sns.barplot(data=df.head(15), x=x, y=y, hue=x, legend=False)
            plt.xticks(rotation=45, ha="right")
        plt.title("PFM 분석 차트")
        plt.tight_layout()
        plt.savefig(path, dpi=150, bbox_inches="tight")
        plt.close()
        return str(path)
    except Exception:
        plt.close()
        return None


# ---------------------------------------------------------------------------
# 종합 통계 & DOCX
# ---------------------------------------------------------------------------
def overall_summary() -> str:
    """상반기 소비 종합 통계 문단을 생성한다."""
    df = load_dataframe()
    total = int(df["Amount_KRW"].sum())
    cat = df.groupby("Category")["Amount_KRW"].sum().sort_values(ascending=False)
    top_cat = cat.index[0] if not cat.empty else "-"
    top_amt = int(cat.iloc[0]) if not cat.empty else 0
    store = df.groupby("Store_Name")["Amount_KRW"].sum().sort_values(ascending=False)
    top_store = store.index[0] if not store.empty else "-"
    lines = [
        f"분석 기간: {df['month'].min()} ~ {df['month'].max()}",
        f"누적 지출액: {total:,}원",
        f"최다 소비 카테고리: {top_cat} ({top_amt:,}원)",
        f"최다 지출 가맹점: {top_store}",
        f"총 결제 건수: {len(df):,}건",
    ]
    return "\n".join(f"- {ln}" for ln in lines)


def generate_docx(entries: list[dict], chart_paths: list[str], trend: str) -> bytes:
    """대화 컨텍스트 + 통계 + 차트로 종합 가계부 진단 보고서를 생성한다."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 표지
    doc.add_heading("개인 맞춤형 자산 관리 가계부 진단 보고서", level=0)
    doc.add_paragraph(f"생성 일시: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph("")

    # 요약 문단
    doc.add_heading("상반기 소비 요약", level=1)
    doc.add_paragraph(overall_summary())

    # 세부 소비 데이터 분석 (대화 컨텍스트)
    doc.add_heading("세부 소비 데이터 분석", level=1)
    for i, e in enumerate(entries, 1):
        doc.add_heading(f"분석 {i}. {e.get('question', '')}", level=2)
        doc.add_paragraph(e.get("answer", ""))

    # 이미지 그래프 섹션
    if chart_paths:
        doc.add_heading("시각화 그래프", level=1)
        for p in chart_paths:
            if Path(p).exists():
                doc.add_picture(p, width=Inches(5.5))
                doc.add_paragraph("")

    # 외부 트렌드 원인 분석
    doc.add_heading("외부 금융 트렌드 분석", level=1)
    doc.add_paragraph(trend or "(웹 검색 정보 없음)")

    # 종합 재무 결론
    doc.add_heading("종합 재무 결론", level=1)
    doc.add_paragraph(
        "상반기 소비 패턴을 종합할 때, 식비/배달과 같은 변동성 큰 카테고리의 "
        "야간 지출을 통제하는 것이 가장 효과적인 절약 전략입니다. "
        "다음 달에는 최다 소비 카테고리 지출을 10% 줄이는 목표를 권장합니다."
    )

    doc.save(str(REPORT_PATH))
    return REPORT_PATH.read_bytes()


# ---------------------------------------------------------------------------
# 세션
# ---------------------------------------------------------------------------
def init_session() -> None:
    defaults = {"messages": [], "report_entries": [], "saved_charts": []}
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def reset_chat() -> None:
    st.session_state.messages = []
    st.session_state.report_entries = []
    st.session_state.saved_charts = []
    st.session_state.pop("docx_bytes", None)


# ---------------------------------------------------------------------------
# 메인
# ---------------------------------------------------------------------------
def main() -> None:
    st.set_page_config(page_title="개인 맞춤형 자산 관리 가계부 Agent", page_icon="💬", layout="wide")
    init_session()
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    st.title("개인 맞춤형 자산 관리 가계부 Agent")
    st.caption("자연어로 소비를 질문하고, 필요할 때 DOCX 보고서를 받으세요.")

    if not CSV_PATH.exists():
        st.error(f"데이터 파일을 찾을 수 없습니다: {CSV_PATH}")
        st.stop()

    api_key = load_api_key()
    if not api_key:
        st.error("OpenAI API Key가 설정되지 않았습니다. .env 파일을 확인해주세요.")
        st.stop()

    try:
        client = OpenAI(api_key=api_key)
    except Exception as exc:  # noqa: BLE001
        st.error(f"OpenAI 클라이언트 초기화 실패: {exc}")
        st.stop()

    # 사이드바
    st.sidebar.header("⚙️ 제어 패널")
    model = st.sidebar.selectbox("LLM 모델 선택", ["gpt-4o", "gpt-4-turbo", "gpt-4o-mini"], index=0)
    use_web = st.sidebar.checkbox("실시간 금융 트렌드 웹 검색 (web_search)", value=False)
    st.sidebar.divider()
    if st.sidebar.button("대화 초기화 (Reset Chat)", width="stretch"):
        reset_chat()
        st.rerun()
    st.sidebar.caption(f"데이터: {len(load_dataframe()):,}건")

    # 지난 대화 렌더링
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            if msg.get("sql"):
                with st.expander("실행된 SQL 보기"):
                    st.code(msg["sql"], language="sql")
            if msg.get("fig") is not None:
                st.plotly_chart(msg["fig"], width="stretch")

    # 입력 처리
    if prompt := st.chat_input("나의 소비 내역이나 가계부에 대해 질문하세요."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        with st.chat_message("assistant"):
            with st.spinner("질문을 SQL로 변환해 분석하는 중..."):
                try:
                    sql = generate_sql(client, model, prompt)
                    df = execute_sql(sql)
                    answer = summarize(client, model, prompt, sql, df)
                    st.write(answer)

                    with st.expander("실행된 SQL 보기"):
                        st.code(sql, language="sql")

                    chart_type = infer_chart_type(df)
                    fig = build_plotly(df, chart_type)
                    if fig is not None:
                        st.plotly_chart(fig, width="stretch")

                    if not df.empty:
                        with st.expander(f"조회 결과 ({len(df)}건)"):
                            st.dataframe(df, width="stretch")

                    # 리포트용 정적 이미지 백업
                    chart_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                    png = save_png(df, chart_type, chart_id)
                    if png:
                        st.session_state.saved_charts.append(png)

                    st.session_state.messages.append(
                        {"role": "assistant", "content": answer, "sql": sql, "fig": fig}
                    )
                    st.session_state.report_entries.append(
                        {"question": prompt, "answer": answer, "sql": sql}
                    )
                except Exception as exc:  # noqa: BLE001
                    err = f"분석 중 오류가 발생했습니다: {exc}"
                    st.write(err)
                    st.session_state.messages.append({"role": "assistant", "content": err})

    # 리포트 생성 영역 (버튼 클릭 시에만 DOCX 생성)
    if st.session_state.report_entries:
        st.divider()
        st.write("위 분석 내용과 시각화 그래프를 기반으로 정식 자산 관리 보고서(DOCX)를 생성해드릴까요?")
        if st.button("📝 자산 관리 보고서(DOCX) 생성"):
            with st.spinner("보고서를 생성하는 중..."):
                try:
                    trend = ""
                    if use_web:
                        trend = run_web_search(client, "가계부 절약 및 소비 트렌드")
                    docx_bytes = generate_docx(
                        st.session_state.report_entries,
                        st.session_state.saved_charts,
                        trend,
                    )
                    st.session_state["docx_bytes"] = docx_bytes
                    st.success("보고서가 생성되었습니다. 아래 버튼으로 다운로드하세요.")
                except Exception as exc:  # noqa: BLE001
                    st.error(f"보고서 생성 오류: {exc}")

        if st.session_state.get("docx_bytes"):
            st.download_button(
                "📥 보고서 다운로드 (DOCX)",
                data=st.session_state["docx_bytes"],
                file_name="PFM_Analysis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


if __name__ == "__main__":
    main()
