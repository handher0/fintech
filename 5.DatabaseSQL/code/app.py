"""
글로벌 핀테크 실시간 AI FDS 관제 및 SQL 에이전트 시스템

아키텍처:
  data/global_fds.db (SQLite 읽기/쓰기)
    → 자연어 질문 → LLM SQL 생성 → 실행(SELECT/INSERT/UPDATE/DELETE 등) → 분석 답변
    → Matplotlib PNG 저장 → python-docx 보고서 생성

실행:
  pip install streamlit plotly matplotlib seaborn python-docx python-dotenv openai pandas
  cd 5.DatabaseSQL/code
  streamlit run app.py
"""

from __future__ import annotations

import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import seaborn as sns
import streamlit as st
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI

# ---------------------------------------------------------------------------
# 경로
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parents[1]
DB_PATH = BASE_DIR / "data" / "global_fds.db"
CHARTS_DIR = BASE_DIR / "output" / "charts"
REPORTS_DIR = BASE_DIR / "output" / "reports"
REPORT_PATH = REPORTS_DIR / "FDS_Analysis_Report.docx"

TABLE_NAME = "global_payments"
SCHEMA_INFO = """
테이블명: global_payments
컬럼:
  - transaction_id (TEXT) : 결제 고유 ID
  - user_id (TEXT) : 고객 ID
  - timestamp (TEXT) : 결제 일시 (YYYY-MM-DD HH:MM:SS)
  - amount_usd (REAL) : 결제 금액 (USD)
  - country (TEXT) : 국가 (US, KR, SG, GB, JP)
  - merchant_category (TEXT) : 업종 (Gaming, Luxury, Food, Electronics, Travel)
  - device_ip (TEXT) : 결제 기기 IP
  - is_fraud (INTEGER) : 0=정상, 1=이상거래
"""

SQL_SYSTEM_PROMPT = f"""당신은 SQLite SQL 전문가입니다. 아래 스키마를 참고해 사용자 요청에 맞는 SQL을 작성하세요.

{SCHEMA_INFO}

규칙:
- SQLite 문법을 사용합니다.
- SELECT, INSERT, UPDATE, DELETE, CREATE, DROP, ALTER 등 필요한 모든 SQL을 사용할 수 있습니다.
- 사용자 의도에 맞는 가장 적절한 SQL 한 개만 작성하세요.
- 응답에는 SQL 쿼리만 출력하고 다른 설명은 하지 마세요.
"""

ANSWER_SYSTEM_PROMPT = """당신은 핀테크 FDS 보안 분석가입니다.
SQL 실행 결과를 바탕으로 한국어로 답변하세요.

답변 형식 규칙 (반드시 준수):
- 일반 문단 텍스트와 불릿 리스트(· 또는 -)만 사용하세요.
- 마크다운 기호(#, **, `, ---)는 절대 사용하지 마세요.
- 숫자는 읽기 쉽게 천 단위 구분을 사용하세요.
- 핵심 인사이트를 2~4개 불릿으로 정리하세요.
- INSERT/UPDATE/DELETE 등 변경 작업이면 영향 받은 행 수와 작업 내용을 명확히 설명하세요.
"""


# ---------------------------------------------------------------------------
# 환경 변수 & DB
# ---------------------------------------------------------------------------
def load_api_key() -> str | None:
    """프로젝트 루트 .env에서 OPENAI_API_KEY를 로드한다."""
    repo_root = BASE_DIR.parent
    load_dotenv(repo_root / ".env", override=False)
    key = os.getenv("OPENAI_API_KEY", "").strip()
    return key or None


def get_db_connection() -> sqlite3.Connection:
    """SQLite 읽기/쓰기 연결을 반환한다."""
    if not DB_PATH.exists():
        raise FileNotFoundError(f"DB 파일을 찾을 수 없습니다: {DB_PATH}")
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def detect_sql_type(sql: str) -> str:
    """SQL 문의 주요 유형을 판별한다."""
    cleaned = re.sub(r"--.*$", "", sql, flags=re.MULTILINE)
    cleaned = re.sub(r"/\*.*?\*/", "", cleaned, flags=re.DOTALL).strip().upper()
    if cleaned.startswith("WITH"):
        return "SELECT"
    for keyword in ("SELECT", "INSERT", "UPDATE", "DELETE", "CREATE", "DROP", "ALTER", "REPLACE"):
        if cleaned.startswith(keyword) or re.match(rf"^\s*{keyword}\b", cleaned):
            return keyword
    return "OTHER"


def extract_sql(text: str) -> str:
    """LLM 응답에서 SQL 쿼리를 추출한다."""
    block = re.search(r"```(?:sql)?\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if block:
        return block.group(1).strip().rstrip(";")

    sql_start = r"(?:WITH|SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|REPLACE)\b"
    match = re.search(rf"(({sql_start}[\s\S]*))", text, re.IGNORECASE)
    if match:
        return match.group(1).strip().rstrip(";")

    return text.strip().rstrip(";")


def execute_sql(sql: str) -> tuple[pd.DataFrame, dict]:
    """SQL을 실행하고 (결과 DataFrame, 실행 정보)를 반환한다."""
    sql_type = detect_sql_type(sql)
    try:
        with get_db_connection() as conn:
            if sql_type == "SELECT":
                df = pd.read_sql_query(sql, conn)
                return df, {"type": "SELECT", "rows_affected": len(df)}

            cursor = conn.execute(sql)
            conn.commit()
            return pd.DataFrame(), {
                "type": sql_type,
                "rows_affected": cursor.rowcount,
            }
    except sqlite3.Error as exc:
        raise RuntimeError(f"SQL 실행 오류: {exc}") from exc


@st.cache_data(show_spinner=False)
def load_kpi_metrics() -> dict[str, float]:
    """상단 KPI 카드용 집계 SQL을 실행한다."""
    sql = f"""
    SELECT
        COALESCE(SUM(amount_usd), 0) AS total_amount,
        COALESCE(SUM(CASE WHEN is_fraud = 1 THEN 1 ELSE 0 END), 0) AS fraud_count,
        COALESCE(SUM(CASE WHEN is_fraud = 1 THEN amount_usd ELSE 0 END), 0) AS blocked_amount
    FROM {TABLE_NAME}
    """
    try:
        with get_db_connection() as conn:
            row = pd.read_sql_query(sql, conn).iloc[0]
        return {
            "total_amount": float(row["total_amount"]),
            "fraud_count": float(row["fraud_count"]),
            "blocked_amount": float(row["blocked_amount"]),
        }
    except Exception:
        return {"total_amount": 0.0, "fraud_count": 0.0, "blocked_amount": 0.0}


# ---------------------------------------------------------------------------
# LLM
# ---------------------------------------------------------------------------
def call_llm(client: OpenAI, model: str, system: str, user: str) -> str:
    """OpenAI Chat Completions API 호출."""
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=0.2,
    )
    return (resp.choices[0].message.content or "").strip()


def generate_sql(client: OpenAI, model: str, question: str) -> str:
    """자연어 질문을 SQL로 변환한다."""
    raw = call_llm(client, model, SQL_SYSTEM_PROMPT, question)
    return extract_sql(raw)


def summarize_results(
    client: OpenAI,
    model: str,
    question: str,
    sql: str,
    df: pd.DataFrame,
    exec_info: dict,
    web_context: str = "",
) -> str:
    """SQL 결과를 일반 텍스트 + 불릿 형태로 요약한다."""
    if exec_info["type"] == "SELECT":
        preview = df.head(50).to_string(index=False) if not df.empty else "(결과 없음)"
        result_block = (
            f"[결과 행 수] {len(df)}건\n"
            f"[결과 데이터 미리보기]\n{preview}"
        )
    else:
        result_block = (
            f"[SQL 유형] {exec_info['type']}\n"
            f"[영향 받은 행 수] {exec_info['rows_affected']}건"
        )

    extra = f"\n\n[웹 검색 참고]\n{web_context}" if web_context else ""
    user_msg = (
        f"[사용자 질문]\n{question}\n\n"
        f"[실행 SQL]\n{sql}\n\n"
        f"{result_block}"
        f"{extra}"
    )
    return call_llm(client, model, ANSWER_SYSTEM_PROMPT, user_msg)


def run_web_search(client: OpenAI, question: str) -> str:
    """OpenAI Responses API web_search로 보조 정보를 수집한다."""
    if not hasattr(client, "responses"):
        return ""
    try:
        resp = client.responses.create(
            model="gpt-5-nano",
            input=f"핀테크 이상거래 분석 맥락에서 다음 질문에 대한 최신 참고 정보를 간략히 정리: {question}",
            tools=[{"type": "web_search"}],
        )
        if hasattr(resp, "output_text") and resp.output_text:
            return str(resp.output_text)
        parts = []
        for item in getattr(resp, "output", []) or []:
            if getattr(item, "type", None) == "message":
                for block in getattr(item, "content", []) or []:
                    if getattr(block, "type", None) == "output_text":
                        parts.append(getattr(block, "text", "") or "")
        return "\n".join(parts).strip()
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# 차트
# ---------------------------------------------------------------------------
def infer_chart_type(df: pd.DataFrame) -> str:
    """결과 DataFrame 성격에 따라 차트 유형을 추론한다."""
    if df.empty or len(df.columns) < 2:
        return "none"

    cols_lower = [c.lower() for c in df.columns]
    time_keywords = ("date", "time", "timestamp", "hour", "day", "month", "year")
    cat_keywords = ("country", "merchant", "category", "user", "업종", "국가")

    if any(any(kw in c for kw in time_keywords) for c in cols_lower):
        return "line"
    if any(any(kw in c for kw in cat_keywords) for c in cols_lower):
        return "bar"

    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if len(numeric_cols) >= 1 and len(df) > 1:
        return "bar"
    return "none"


def build_plotly_chart(df: pd.DataFrame, chart_type: str) -> go.Figure | None:
    """Plotly 인터랙티브 차트를 생성한다."""
    if chart_type == "none" or df.empty:
        return None

    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if not numeric_cols:
        return None

    y_col = numeric_cols[0]
    x_col = df.columns[0] if df.columns[0] != y_col else (df.columns[1] if len(df.columns) > 1 else df.columns[0])

    if chart_type == "line":
        fig = px.line(df, x=x_col, y=y_col, title="시간별/추이 분석", markers=True)
    else:
        fig = px.bar(df, x=x_col, y=y_col, title="국가/업종별 비교 분석", color=x_col)

    fig.update_layout(hovermode="x unified", showlegend=False)
    return fig


def save_static_chart_png(df: pd.DataFrame, chart_type: str, chart_id: str) -> str | None:
    """Matplotlib/Seaborn으로 정적 PNG를 output/charts/에 저장한다."""
    if chart_type == "none" or df.empty:
        return None

    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    filepath = CHARTS_DIR / f"chart_{chart_id}.png"

    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if not numeric_cols:
        return None

    y_col = numeric_cols[0]
    x_col = df.columns[0] if df.columns[0] != y_col else df.columns[1]

    try:
        plt.figure(figsize=(10, 5))
        sns.set_style("whitegrid")
        if chart_type == "line":
            sns.lineplot(data=df, x=x_col, y=y_col, marker="o")
            plt.title("시간별 추이 분석")
        else:
            sns.barplot(data=df.head(20), x=x_col, y=y_col, hue=x_col, legend=False)
            plt.title("비교 분석 차트")
            plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        plt.savefig(filepath, dpi=150, bbox_inches="tight")
        plt.close()
        return str(filepath)
    except Exception:
        plt.close()
        return None


# ---------------------------------------------------------------------------
# 보고서 (DOCX)
# ---------------------------------------------------------------------------
def generate_docx_report(entries: list[dict], chart_paths: list[str]) -> bytes:
    """대화 내역과 차트 이미지를 Word 보고서로 컴파일한다."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("글로벌 핀테크 FDS 분석 보고서", level=0)
    doc.add_paragraph(f"생성 일시: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph("본 보고서는 AI SQL Agent 대화 분석 결과를 자동 생성한 문서입니다.")
    doc.add_paragraph("")

    for i, entry in enumerate(entries, 1):
        doc.add_heading(f"분석 {i}", level=1)
        doc.add_heading("사용자 질문", level=2)
        doc.add_paragraph(entry.get("question", ""))
        doc.add_heading("실행 SQL", level=2)
        doc.add_paragraph(entry.get("sql", "(없음)"))
        doc.add_heading("분석 결과", level=2)
        doc.add_paragraph(entry.get("answer", ""))
        doc.add_paragraph("")

    if chart_paths:
        doc.add_heading("시각화 차트", level=1)
        for path in chart_paths:
            if Path(path).exists():
                doc.add_paragraph(Path(path).name)
                doc.add_picture(path, width=Inches(5.5))
                doc.add_paragraph("")

    doc.save(str(REPORT_PATH))
    return REPORT_PATH.read_bytes()


# ---------------------------------------------------------------------------
# 세션 초기화
# ---------------------------------------------------------------------------
def init_session() -> None:
    defaults = {
        "messages": [],       # {role, content}
        "report_entries": [], # {question, answer, sql}
        "saved_charts": [],   # png path list
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def clear_chat() -> None:
    st.session_state.messages = []
    st.session_state.report_entries = []
    st.session_state.saved_charts = []


# ---------------------------------------------------------------------------
# SQL Agent 파이프라인
# ---------------------------------------------------------------------------
def process_user_question(
    question: str,
    client: OpenAI,
    model: str,
    use_web_search: bool,
) -> tuple[str, str, pd.DataFrame, dict, object, str | None]:
    """질문 → SQL → 실행 → 요약 → 차트 생성 파이프라인."""
    sql = generate_sql(client, model, question)
    df, exec_info = execute_sql(sql)

    # 데이터 변경 후 KPI 캐시 갱신
    if exec_info["type"] != "SELECT":
        load_kpi_metrics.clear()

    web_context = run_web_search(client, question) if use_web_search else ""
    answer = summarize_results(client, model, question, sql, df, exec_info, web_context)

    fig, png_path = None, None
    if exec_info["type"] == "SELECT" and not df.empty:
        chart_type = infer_chart_type(df)
        chart_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        png_path = save_static_chart_png(df, chart_type, chart_id)
        fig = build_plotly_chart(df, chart_type)
        if png_path:
            st.session_state.saved_charts.append(png_path)

    return answer, sql, df, exec_info, fig, png_path


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------
def render_kpi() -> None:
    """상단 KPI 메트릭 카드를 렌더링한다."""
    kpi = load_kpi_metrics()
    c1, c2, c3 = st.columns(3)
    c1.metric("총 모니터링 금액", f"${kpi['total_amount']:,.2f}")
    c2.metric("탐지된 이상거래 건수", f"{int(kpi['fraud_count']):,}건")
    c3.metric("피해 예방 차단 금액", f"${kpi['blocked_amount']:,.2f}")


def render_sidebar() -> tuple[str, bool]:
    """사이드바 제어 패널."""
    st.sidebar.header("⚙️ 제어 패널")
    model = st.sidebar.selectbox(
        "LLM 모델 선택",
        ["gpt-4o", "gpt-4-turbo", "gpt-4o-mini"],
        index=0,
    )
    use_web = st.sidebar.checkbox("OpenAI 실시간 웹 검색 (web_search)", value=False)
    st.sidebar.info("전체 SQL(SELECT/INSERT/UPDATE/DELETE/CREATE 등) 실행 가능")
    st.sidebar.divider()
    if st.sidebar.button("대화 초기화 (Clear Chat)", width="stretch"):
        clear_chat()
        st.rerun()
    st.sidebar.caption(f"DB: {DB_PATH.name}")
    st.sidebar.caption(f"저장된 차트: {len(st.session_state.saved_charts)}개")
    return model, use_web


def main() -> None:
    st.set_page_config(
        page_title="AI FDS SQL Agent",
        page_icon="🛡️",
        layout="wide",
    )
    init_session()
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    st.title("🛡️ 글로벌 핀테크 실시간 AI FDS 관제 및 SQL 에이전트 시스템")
    st.caption("자연어 질문 → SQL 자동 생성 → 이상거래 분석 · 시각화 · 보고서")

    # API Key 확인
    api_key = load_api_key()
    if not api_key:
        st.error("OpenAI API Key가 설정되지 않았습니다. .env 파일을 확인해주세요.")
        st.stop()

    if not DB_PATH.exists():
        st.error(f"데이터베이스를 찾을 수 없습니다: {DB_PATH}")
        st.stop()

    try:
        client = OpenAI(api_key=api_key)
    except Exception as exc:
        st.error(f"OpenAI 클라이언트 초기화 실패: {exc}")
        st.stop()

    model, use_web = render_sidebar()

    # KPI
    render_kpi()
    st.divider()

    # 보고서 생성 버튼 (사이드바 하단)
    with st.sidebar:
        st.divider()
        if st.button("📝 분석 보고서 자동 생성", width="stretch"):
            if not st.session_state.report_entries:
                st.sidebar.warning("아직 분석 대화가 없습니다.")
            else:
                try:
                    docx_bytes = generate_docx_report(
                        st.session_state.report_entries,
                        st.session_state.saved_charts,
                    )
                    st.session_state["docx_bytes"] = docx_bytes
                    st.sidebar.success("보고서가 생성되었습니다.")
                except Exception as exc:
                    st.sidebar.error(f"보고서 생성 오류: {exc}")

        if st.session_state.get("docx_bytes"):
            st.download_button(
                label="📥 DOCX 다운로드",
                data=st.session_state["docx_bytes"],
                file_name="FDS_Analysis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width="stretch",
            )

    # 채팅 히스토리 렌더링
    st.subheader("💬 SQL Agent 대화형 분석")
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            if msg.get("sql"):
                with st.expander("실행된 SQL 보기"):
                    st.code(msg["sql"], language="sql")
            if msg.get("fig"):
                st.plotly_chart(msg["fig"], width="stretch")

    # 사용자 입력
    if prompt := st.chat_input("질문을 입력하세요 (조회·수정·삭제·테이블 생성 등 모두 가능)"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        with st.chat_message("assistant"):
            with st.spinner("SQL을 생성하고 분석하는 중..."):
                try:
                    answer, sql, df, exec_info, fig, png_path = process_user_question(
                        prompt, client, model, use_web
                    )
                    st.write(answer)
                    with st.expander("실행된 SQL 보기"):
                        st.code(sql, language="sql")

                    if fig is not None:
                        st.plotly_chart(fig, width="stretch")

                    if exec_info.get("type") == "SELECT" and df is not None:
                        if not df.empty:
                            with st.expander(f"조회 결과 ({len(df)}건)"):
                                st.dataframe(df, width="stretch")
                        else:
                            st.caption("조회 결과가 없습니다.")

                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": answer,
                        "sql": sql,
                        "fig": fig,
                    })
                    st.session_state.report_entries.append({
                        "question": prompt,
                        "answer": answer,
                        "sql": sql,
                    })
                except Exception as exc:
                    err = f"분석 중 오류가 발생했습니다: {exc}"
                    st.write(err)
                    st.session_state.messages.append({"role": "assistant", "content": err})


if __name__ == "__main__":
    main()
